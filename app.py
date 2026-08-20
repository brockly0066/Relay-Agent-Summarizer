import json
import time
import tempfile
import os
import io
import base64
import streamlit as st
from google import genai
from google.genai import types
from docx import Document
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build

TEXT_EXTS = {"txt", "eml", "md"}
DOCX_EXTS = {"docx"}
BINARY_EXTS = {"pdf", "png", "jpg", "jpeg"}


def extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ── Gmail OAuth ──────────────────────────────────────────────────────────
GMAIL_SCOPES = ["https://www.googleapis.com/auth/gmail.readonly"]


def get_gmail_oauth_config():
    try:
        return {
            "client_id": st.secrets["GOOGLE_CLIENT_ID"],
            "client_secret": st.secrets["GOOGLE_CLIENT_SECRET"],
            "redirect_uri": st.secrets["GOOGLE_REDIRECT_URI"],
        }
    except Exception:
        return None


def build_flow(cfg):
    return Flow.from_client_config(
        {
            "web": {
                "client_id": cfg["client_id"],
                "client_secret": cfg["client_secret"],
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "redirect_uris": [cfg["redirect_uri"]],
            }
        },
        scopes=GMAIL_SCOPES,
        redirect_uri=cfg["redirect_uri"],
    )


def gmail_service_from_session():
    creds_dict = st.session_state.get("_gmail_creds")
    if not creds_dict:
        return None
    from google.oauth2.credentials import Credentials
    creds = Credentials(**creds_dict)
    return build("gmail", "v1", credentials=creds)


def decode_body(data: str) -> str:
    if not data:
        return ""
    padded = data + "=" * (-len(data) % 4)
    return base64.urlsafe_b64decode(padded).decode("utf-8", errors="ignore")


def strip_html(html: str) -> str:
    import re
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</p>", "\n\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()


def extract_body_from_payload(payload) -> str:
    if not payload:
        return ""
    mime_type = payload.get("mimeType", "")
    body = payload.get("body", {})
    if mime_type == "text/plain" and body.get("data"):
        return decode_body(body["data"])
    if mime_type == "text/html" and body.get("data"):
        return strip_html(decode_body(body["data"]))
    plain, html = "", ""
    for part in payload.get("parts", []) or []:
        pm = part.get("mimeType", "")
        pb = part.get("body", {})
        if pm == "text/plain" and pb.get("data") and not plain:
            plain = decode_body(pb["data"])
        elif pm == "text/html" and pb.get("data") and not html:
            html = strip_html(decode_body(pb["data"]))
        elif part.get("parts"):
            nested = extract_body_from_payload(part)
            if nested and not plain:
                plain = nested
    return plain or html


def header_value(headers, name):
    for h in headers or []:
        if h.get("name", "").lower() == name.lower():
            return h.get("value", "")
    return ""


def format_thread_as_text(thread: dict) -> str:
    blocks = []
    for msg in thread.get("messages", []):
        headers = msg.get("payload", {}).get("headers", [])
        sender = header_value(headers, "From")
        date = header_value(headers, "Date")
        subject = header_value(headers, "Subject")
        body = extract_body_from_payload(msg.get("payload", {}))
        blocks.append(f"From: {sender}\nDate: {date}\nSubject: {subject}\n\n{body.strip()}")
    return "\n\n---\n\n".join(blocks)

# ── Page setup ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Relay — Agent Trace Summarizer",
    page_icon="📡",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Theme / CSS ──────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,400;9..144,500;9..144,600;9..144,700&family=Inter:wght@400;500;600;700;800&family=JetBrains+Mono:wght@400;500;600&display=swap');

:root {
    --void: #100B1C;
    --void-2: #170F28;
    --glass: rgba(255,255,255,0.045);
    --glass-raised: rgba(255,255,255,0.07);
    --hairline: rgba(255,255,255,0.10);
    --text: #F3EFE8;
    --text-dim: #A79FC0;
    --coral: #FF7A5C;
    --coral-soft: rgba(255,122,92,0.16);
    --cyan: #6EE7D8;
    --cyan-soft: rgba(110,231,216,0.16);
    --violet: #8B7FE8;
    --red: #FF6B7A;
}

html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

/* Ambient gradient-mesh backdrop — the signature element */
.stApp {
    background-color: var(--void);
    background-image:
        radial-gradient(ellipse 900px 700px at 8% -5%, rgba(255,122,92,0.20), transparent 60%),
        radial-gradient(ellipse 800px 800px at 105% 10%, rgba(139,127,232,0.22), transparent 55%),
        radial-gradient(ellipse 700px 600px at 50% 115%, rgba(110,231,216,0.14), transparent 55%),
        linear-gradient(180deg, var(--void) 0%, var(--void-2) 100%);
    background-attachment: fixed;
}

#MainMenu, footer, header { visibility: hidden; }

/* ── Header ── */
.eyebrow {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11.5px;
    letter-spacing: 0.16em;
    text-transform: uppercase;
    color: var(--coral);
    display: flex;
    align-items: center;
    gap: 9px;
    margin-bottom: 18px;
}
.eyebrow .dot {
    width: 7px; height: 7px;
    background: var(--coral);
    border-radius: 50%;
    box-shadow: 0 0 12px var(--coral);
    animation: breathe 2.2s ease-in-out infinite;
}
@keyframes breathe { 0%,100% { opacity: 1; } 50% { opacity: 0.4; } }

.hero-title {
    font-family: 'Fraunces', serif;
    font-optical-sizing: auto;
    font-size: clamp(42px, 5.5vw, 68px);
    font-weight: 600;
    line-height: 1.02;
    letter-spacing: -0.02em;
    margin: 0 0 20px 0;
    color: var(--text);
}
.hero-title .accent {
    background: linear-gradient(100deg, var(--coral), var(--violet) 60%, var(--cyan));
    -webkit-background-clip: text;
    background-clip: text;
    -webkit-text-fill-color: transparent;
}
.hero-title .rest {
    color: var(--text-dim);
    font-weight: 400;
    font-style: italic;
    font-size: 0.55em;
    display: block;
    margin-top: 6px;
    letter-spacing: 0;
}
.hero-sub {
    color: var(--text-dim);
    font-size: 16px;
    line-height: 1.7;
    max-width: 620px;
    margin-bottom: 6px;
}

/* connection status pill */
.status-pill {
    display: inline-flex;
    align-items: center;
    gap: 8px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 12px;
    padding: 7px 14px;
    border-radius: 100px;
    margin-top: 20px;
}
.status-pill.ok { background: var(--cyan-soft); color: var(--cyan); border: 1px solid rgba(110,231,216,0.3); }
.status-pill.warn { background: rgba(255,122,92,0.14); color: var(--coral); border: 1px solid rgba(255,122,92,0.3); }
.status-pill .sdot { width: 6px; height: 6px; border-radius: 50%; background: currentColor; }

/* panel label */
.panel-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: var(--text-dim);
    margin-bottom: 6px;
    display: block;
}
.panel-num { color: var(--coral); }

/* Glass panels via Streamlit's bordered container */
div[data-testid="stVerticalBlockBorderWrapper"] {
    background: var(--glass) !important;
    backdrop-filter: blur(24px);
    -webkit-backdrop-filter: blur(24px);
    border: 1px solid var(--hairline) !important;
    border-radius: 18px !important;
    box-shadow: 0 8px 32px rgba(0,0,0,0.25);
}

/* Inputs */
.stTextArea textarea, .stSelectbox div[data-baseweb="select"] > div {
    background: rgba(255,255,255,0.03) !important;
    border: 1px solid var(--hairline) !important;
    color: var(--text) !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 13.5px !important;
    border-radius: 10px !important;
}
.stTextArea textarea:focus { border-color: var(--coral) !important; box-shadow: 0 0 0 1px var(--coral) !important; }
label, .stCaption, [data-testid="stCaptionContainer"] { color: var(--text-dim) !important; }
label { font-size: 12.5px !important; font-weight: 500 !important; }

/* Radio pills */
div[role="radiogroup"] { gap: 8px; }
div[role="radiogroup"] label {
    background: rgba(255,255,255,0.04);
    border: 1px solid var(--hairline);
    border-radius: 100px;
    padding: 7px 16px !important;
    transition: all 0.15s;
}

/* Buttons */
.stButton > button {
    background: linear-gradient(100deg, var(--coral), #E85A6A) !important;
    color: #1A0F0C !important;
    border: none !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 700 !important;
    font-size: 14.5px !important;
    border-radius: 12px !important;
    padding: 12px 20px !important;
    width: 100%;
    box-shadow: 0 4px 20px rgba(255,122,92,0.25);
    transition: transform 0.12s, box-shadow 0.15s;
}
.stButton > button:hover { box-shadow: 0 6px 28px rgba(255,122,92,0.4) !important; transform: translateY(-1px); color: #1A0F0C !important; }
.stButton > button:disabled { opacity: 0.4 !important; box-shadow: none !important; }

/* Gmail connect button */
.gmail-connect-btn {
    display: inline-flex;
    align-items: center;
    gap: 10px;
    background: rgba(255,255,255,0.05);
    border: 1px solid var(--hairline);
    color: var(--text) !important;
    font-family: 'Inter', sans-serif;
    font-weight: 600;
    font-size: 14px;
    border-radius: 12px;
    padding: 11px 20px;
    text-decoration: none !important;
    transition: all 0.15s;
}
.gmail-connect-btn:hover {
    border-color: var(--coral);
    background: rgba(255,122,92,0.08);
    transform: translateY(-1px);
}
.gmail-connect-btn svg { flex-shrink: 0; }

/* File uploader */
[data-testid="stFileUploaderDropzone"] {
    background: rgba(255,255,255,0.03) !important;
    border: 1px dashed var(--hairline) !important;
    border-radius: 12px !important;
}

/* Status widget (agent trace) */
[data-testid="stStatusWidget"], div[data-testid="stExpander"] {
    background: var(--glass) !important;
    backdrop-filter: blur(24px);
    border: 1px solid var(--hairline) !important;
    border-radius: 16px !important;
}
[data-testid="stExpander"] summary, [data-testid="stExpander"] p { font-family: 'JetBrains Mono', monospace !important; font-size: 13px !important; }

/* Output */
.summary-card { font-size: 15px; line-height: 1.75; color: var(--text); }

.action-item {
    display: grid;
    grid-template-columns: 1fr auto auto;
    gap: 16px;
    align-items: start;
    background: rgba(255,255,255,0.035);
    border: 1px solid var(--hairline);
    border-left: 3px solid var(--coral);
    border-radius: 12px;
    padding: 14px 16px;
    margin-bottom: 10px;
}
.action-item .task { font-size: 14px; line-height: 1.55; color: var(--text); }
.action-item .meta { font-family: 'JetBrains Mono', monospace; font-size: 11px; color: var(--text-dim); white-space: nowrap; text-align: right; }
.action-item .owner { color: var(--cyan); font-weight: 600; }
.badge {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    padding: 4px 10px;
    border-radius: 100px;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    white-space: nowrap;
    height: fit-content;
}
.badge.high { background: rgba(255,107,122,0.16); color: var(--red); }
.badge.medium { background: var(--coral-soft); color: var(--coral); }
.badge.low { background: var(--cyan-soft); color: var(--cyan); }

.decision-row {
    font-size: 14px;
    line-height: 1.65;
    padding: 11px 0;
    border-bottom: 1px solid var(--hairline);
    color: var(--text);
}
.decision-row:last-child { border-bottom: none; }
.decision-row::before { content: '◆ '; color: var(--violet); font-size: 10px; }

.empty-note { color: var(--text-dim); font-size: 13px; font-style: italic; }

.footnote {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11.5px;
    color: #5C5478;
    line-height: 1.8;
    margin-top: 34px;
}
.footnote a { color: var(--cyan); }
</style>
""", unsafe_allow_html=True)

# ── API key from Streamlit secrets ──────────────────────────────────────
def get_api_key():
    try:
        return st.secrets["GEMINI_API_KEY"]
    except Exception:
        return None

api_key = get_api_key()
gmail_cfg = get_gmail_oauth_config()

# Handle the redirect back from Google's consent screen
if gmail_cfg and "code" in st.query_params and "_gmail_creds" not in st.session_state:
    try:
        flow = build_flow(gmail_cfg)
        flow.fetch_token(code=st.query_params["code"])
        creds = flow.credentials
        st.session_state["_gmail_creds"] = {
            "token": creds.token,
            "refresh_token": creds.refresh_token,
            "token_uri": creds.token_uri,
            "client_id": creds.client_id,
            "client_secret": creds.client_secret,
            "scopes": creds.scopes,
        }
        st.query_params.clear()
        st.rerun()
    except Exception as e:
        st.error(f"Gmail connection failed: {e}")

# ── Header ───────────────────────────────────────────────────────────────
st.markdown("""
<div class="eyebrow"><span class="dot"></span>Agent online</div>
<div class="hero-title"><span class="accent">Relay</span><span class="rest">thread &amp; transcript triage agent</span></div>
<div class="hero-sub">Drop in an email thread, meeting transcript, PDF, or screenshot. The agent parses it, pulls out
who owns what, and hands back a clean summary — no more re-reading a 40-message thread to find the one commitment
that mattered.</div>
""", unsafe_allow_html=True)

if api_key:
    st.markdown('<div class="status-pill ok"><span class="sdot"></span>Connected to Gemini via Streamlit secrets</div>', unsafe_allow_html=True)
else:
    st.markdown('<div class="status-pill warn"><span class="sdot"></span>No API key configured yet — see setup below</div>', unsafe_allow_html=True)

st.write("")
st.write("")

model_name = "gemini-2.5-flash"

# ── Source panel ─────────────────────────────────────────────────────────
with st.container(border=True):
    st.markdown('<span class="panel-label"><span class="panel-num">01</span> — Source</span>', unsafe_allow_html=True)
    source_type = st.radio("Source type", ["Email thread", "Meeting transcript"],
                            horizontal=True, label_visibility="collapsed")

    gmail_service = gmail_service_from_session() if source_type == "Email thread" else None
    input_method = "Paste or upload"
    if source_type == "Email thread" and gmail_cfg:
        input_method = st.radio("Input method", ["Paste or upload", "Connect Gmail"],
                                 horizontal=True, label_visibility="collapsed")

    pasted_text = ""

    if input_method == "Connect Gmail":
        if not gmail_service:
            flow = build_flow(gmail_cfg)
            auth_url, _ = flow.authorization_url(access_type="offline", prompt="consent")
            st.markdown(f"""
            <a href="{auth_url}" target="_self" class="gmail-connect-btn">
                <svg width="18" height="18" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
                    <path d="M22 5.5v13c0 1.1-.9 2-2 2h-2.5V10.7l-5.5 4.6-5.5-4.6v9.8H4c-1.1 0-2-.9-2-2v-13c0-.85.53-1.58 1.28-1.87L12 10.5l8.72-6.87c.75.29 1.28 1.02 1.28 1.87z" fill="#FF7A5C"/>
                    <path d="M2 5.5c0-.85.53-1.58 1.28-1.87L12 10.5V21H4c-1.1 0-2-.9-2-2v-13.5z" fill="#8B7FE8" fill-opacity="0.55"/>
                    <path d="M22 5.5c0-.85-.53-1.58-1.28-1.87L12 10.5V21h8c1.1 0 2-.9 2-2v-13.5z" fill="#6EE7D8" fill-opacity="0.45"/>
                </svg>
                <span>Connect Gmail</span>
            </a>
            """, unsafe_allow_html=True)
            st.caption("Opens Google's consent screen. You'll be redirected back here once you approve access.")
        else:
            col_a, col_b = st.columns([4, 1])
            with col_a:
                gmail_query = st.text_input("Search Gmail", placeholder="from:alex subject:budget",
                                             label_visibility="collapsed")
            with col_b:
                search_clicked = st.button("Search")
            if search_clicked or "_gmail_results" in st.session_state:
                if search_clicked:
                    try:
                        resp = gmail_service.users().messages().list(
                            userId="me", q=gmail_query or "", maxResults=10
                        ).execute()
                        results = []
                        for m in resp.get("messages", []):
                            meta = gmail_service.users().messages().get(
                                userId="me", id=m["id"], format="metadata",
                                metadataHeaders=["Subject", "From", "Date"],
                            ).execute()
                            headers = meta.get("payload", {}).get("headers", [])
                            results.append({
                                "thread_id": meta["threadId"],
                                "subject": header_value(headers, "Subject") or "(no subject)",
                                "from": header_value(headers, "From"),
                                "date": header_value(headers, "Date"),
                            })
                        st.session_state["_gmail_results"] = results
                    except Exception as e:
                        st.error(f"Gmail search failed: {e}")
                        st.session_state["_gmail_results"] = []

                results = st.session_state.get("_gmail_results", [])
                if results:
                    labels = [f"{r['subject']} — {r['from']} ({r['date']})" for r in results]
                    picked = st.selectbox("Pick a thread", labels, label_visibility="collapsed")
                    picked_idx = labels.index(picked)
                    if st.button("Load this thread"):
                        try:
                            thread_id = results[picked_idx]["thread_id"]
                            thread = gmail_service.users().threads().get(
                                userId="me", id=thread_id, format="full"
                            ).execute()
                            st.session_state["_loaded_gmail_text"] = format_thread_as_text(thread)
                        except Exception as e:
                            st.error(f"Couldn't load thread: {e}")
                elif search_clicked:
                    st.caption("No matching threads found.")

            if st.session_state.get("_loaded_gmail_text"):
                st.text_area("Loaded thread", value=st.session_state["_loaded_gmail_text"],
                              height=220, disabled=True)
                pasted_text = st.session_state["_loaded_gmail_text"]
                if st.button("Disconnect Gmail"):
                    st.session_state.pop("_gmail_creds", None)
                    st.session_state.pop("_gmail_results", None)
                    st.session_state.pop("_loaded_gmail_text", None)
                    st.rerun()

    if input_method == "Paste or upload":
        placeholder = ("Paste the full email thread here — include sender names and timestamps if you have them. "
                        "The agent reads the whole thing, not just the latest reply.") if source_type == "Email thread" \
            else "Paste the meeting transcript here — speaker labels help, but plain text works too."

        pasted_text = st.text_area("Content", height=220, placeholder=placeholder, label_visibility="collapsed")

        uploaded = st.file_uploader(
            "or drop in a file — .txt, .eml, .md, .docx, .pdf, or a screenshot (.png / .jpg)",
            type=["txt", "eml", "md", "docx", "pdf", "png", "jpg", "jpeg"],
        )

        if uploaded is not None:
            ext = uploaded.name.rsplit(".", 1)[-1].lower()
            if ext in TEXT_EXTS:
                pasted_text = uploaded.read().decode("utf-8", errors="ignore")
                st.caption(f"Loaded text from **{uploaded.name}**.")
            elif ext in DOCX_EXTS:
                try:
                    pasted_text = extract_docx_text(uploaded.getvalue())
                    st.caption(f"Extracted text from **{uploaded.name}**.")
                except Exception as e:
                    st.error(f"Couldn't read {uploaded.name}: {e}")
            elif ext in BINARY_EXTS:
                st.caption(f"**{uploaded.name}** will be read directly by Gemini — no paste needed.")
                st.session_state["_pending_upload"] = {
                    "bytes": uploaded.getvalue(),
                    "suffix": f".{ext}",
                    "mime": {
                        "pdf": "application/pdf",
                        "png": "image/png",
                        "jpg": "image/jpeg",
                        "jpeg": "image/jpeg",
                    }[ext],
                    "name": uploaded.name,
                }
        else:
            st.session_state.pop("_pending_upload", None)

    run_clicked = st.button("Run agent", disabled=not api_key)
    if not api_key:
        st.caption("Add `GEMINI_API_KEY` in Streamlit secrets to enable this.")

st.write("")

# ── Prompt builder ───────────────────────────────────────────────────────
def build_prompt(kind: str, text: str = None, has_attachment: bool = False) -> str:
    label = "email thread" if kind == "Email thread" else "meeting transcript"
    source_desc = (
        "attached file (it may be a PDF document or a screenshot — read any visible text and layout) "
        f"containing the {label}" if has_attachment else f"following {label}"
    )
    instructions = f"""You are an operations agent that reads {label}s and extracts what matters.

Read the {source_desc} and respond with ONLY a JSON object (no markdown fences, no preamble) matching this exact shape:

{{
  "summary": "a tight 3-5 sentence summary of what was discussed and decided",
  "action_items": [
    {{ "task": "specific task description", "owner": "person responsible, or 'Unassigned' if unclear", "due_date": "date or timeframe if mentioned, else 'Not specified'", "priority": "high" | "medium" | "low" }}
  ],
  "key_decisions": ["decision 1", "decision 2"]
}}

If there are no clear action items or decisions, return empty arrays for those fields — do not invent content that isn't in the source."""

    if has_attachment:
        return instructions
    return f"""{instructions}

SOURCE:
\"\"\"
{text}
\"\"\""""

STAGES = {
    "Email thread": [
        "Ingesting thread",
        "Resolving participants",
        "Tracing reply chain",
        "Extracting commitments",
        "Structuring output",
    ],
    "Meeting transcript": [
        "Ingesting transcript",
        "Identifying speakers",
        "Segmenting discussion",
        "Extracting commitments",
        "Structuring output",
    ],
}

# ── Run ──────────────────────────────────────────────────────────────────
pending_upload = st.session_state.get("_pending_upload")
has_text = bool(pasted_text and pasted_text.strip())
has_attachment = pending_upload is not None

if run_clicked:
    if not has_text and not has_attachment:
        st.error("Paste an email thread or transcript, or upload a file, first.")
    else:
        stages = STAGES[source_type]
        if has_attachment:
            stages = [f"Reading {pending_upload['name']}"] + stages[1:]

        result = None
        error_msg = None
        tmp_path = None

        with st.status("Running agent…", expanded=True) as status:
            for stage in stages[:-1]:
                st.write(f"◆ {stage}")
                time.sleep(0.35)
            st.write(f"◆ {stages[-1]}")

            try:
                client = genai.Client(api_key=api_key)

                if has_attachment:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=pending_upload["suffix"]) as tmp:
                        tmp.write(pending_upload["bytes"])
                        tmp_path = tmp.name
                    gemini_file = client.files.upload(
                        file=tmp_path,
                        config=types.UploadFileConfig(mime_type=pending_upload["mime"]),
                    )
                    contents = [gemini_file, build_prompt(source_type, has_attachment=True)]
                else:
                    contents = build_prompt(source_type, text=pasted_text)

                response = client.models.generate_content(
                    model=model_name,
                    contents=contents,
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        temperature=0.2,
                    ),
                )
                raw = response.text.strip().replace("```json", "").replace("```", "").strip()
                result = json.loads(raw)
                status.update(label="Agent finished", state="complete", expanded=False)
            except Exception as e:
                error_msg = str(e)
                status.update(label="Agent failed", state="error", expanded=True)
            finally:
                if tmp_path and os.path.exists(tmp_path):
                    os.remove(tmp_path)

        if error_msg:
            st.error(f"Something went wrong: {error_msg}")
        elif result:
            with st.container(border=True):
                st.markdown('<span class="panel-label">Summary</span>', unsafe_allow_html=True)
                st.markdown(f'<div class="summary-card">{result.get("summary", "No summary returned.")}</div>',
                            unsafe_allow_html=True)

            st.write("")
            with st.container(border=True):
                st.markdown('<span class="panel-label">Action items</span>', unsafe_allow_html=True)
                items = result.get("action_items", [])
                if items:
                    for it in items:
                        priority = (it.get("priority") or "medium").lower()
                        st.markdown(f"""
                        <div class="action-item">
                            <div class="task">{it.get('task', '')}</div>
                            <div class="meta"><span class="owner">{it.get('owner', 'Unassigned')}</span><br>{it.get('due_date', 'Not specified')}</div>
                            <div class="badge {priority}">{priority}</div>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.markdown('<div class="empty-note">No action items detected.</div>', unsafe_allow_html=True)

            st.write("")
            with st.container(border=True):
                st.markdown('<span class="panel-label">Key decisions</span>', unsafe_allow_html=True)
                decisions = result.get("key_decisions", [])
                if decisions:
                    for d in decisions:
                        st.markdown(f'<div class="decision-row">{d}</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="empty-note">No explicit decisions detected.</div>', unsafe_allow_html=True)

st.markdown("""
<div class="footnote">
Relay v0.3 · reads its Gemini key from Streamlit secrets, never from the page<br>
"Connect Gmail" needs GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET / GOOGLE_REDIRECT_URI in secrets — see README.
</div>
""", unsafe_allow_html=True)
